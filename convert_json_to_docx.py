import subprocess
import json
import argparse
import tempfile
import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

# Define popular fonts for help text
POPULAR_FONTS = [
    "Arial", "Cambria", "Courier New", "Garamond",
    "Georgia", "Calibri", "Segoe UI", "Times New Roman","Verdana"
]

# Setup argument parser
parser = argparse.ArgumentParser(description="Extract command output and export it as a DOCX table.")
parser.add_argument("targetpath", nargs="?", default="templates.docx", help="Optional path to save the DOCX file (default: current folder)")
parser.add_argument("--command", default="dotnet new list", help="Command to run instead of `dotnet new list`")
parser.add_argument("--font", default="Calibri", help=f"Font for DOCX table (default: 'Calibri'). Popular fonts: {', '.join(POPULAR_FONTS)}")
args = parser.parse_args()

# Create a temporary file for JSON
temp_json_file = os.path.join(tempfile.gettempdir(), "command_output.json")

# Step 1: Extract column headers dynamically
print(f"Running command: {args.command}")
column_headers_command = [
    "powershell", "-Command",
    f"({args.command}) | Select-String -Pattern 'Name\\s+Short Name' | ForEach-Object {{ ($_ -split '\\s{{2,}}') }}"
]
headers_process = subprocess.run(column_headers_command, capture_output=True, text=True, shell=True)
column_headers = headers_process.stdout.strip().split("\n")
print(f"Extracted column headers: {column_headers}")
if not column_headers:
    print("Error: Unable to extract column headers.")
    exit(1)

# Step 2: Extract all table rows properly
print("Extracting table rows...")
rows_command = [
    "powershell", "-Command",
    f"({args.command}) | Select-String -Pattern '^\\S' | Select-Object -Skip 2"
]
rows_process = subprocess.run(rows_command, capture_output=True, text=True, shell=True)
table_rows = rows_process.stdout.strip().split("\n")

if not table_rows or table_rows == [""]:
    print("Error: No data rows found. Ensure the command produces output.")
    exit(1)

# Step 3: Process rows into structured JSON
print("Processing table data into structured JSON...")
data = []
for row in table_rows:
    values = re.split(r'\s{2,}', row.strip())  # Splitting by two or more spaces
    if len(values) == len(column_headers):  # Ensure row matches column count
        entry = dict(zip(column_headers, values))
        data.append(entry)

# Step 4: Save extracted data as JSON
with open(temp_json_file, "w", encoding="utf-8") as f:
    json.dump(data, f, indent=4)

print(f"Temporary JSON file created: {temp_json_file}")

# Step 5: Convert JSON to DOCX (Using a Table)
print("Generating DOCX file...")
doc = Document()
doc.add_heading(f"{args.command} Output Table", level=1)

# Create a table in DOCX
table = doc.add_table(rows=1, cols=len(column_headers))
table.style = "Table Grid"

# Function to set font style for a table cell
def set_font(cell, font_name, font_size=11, bold=False):
    """Applies a custom font to a DOCX table cell."""
    run = cell.paragraphs[0].add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

# Add headers to the table
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(column_headers):
    hdr_cells[i].text = column_name
    set_font(hdr_cells[i], args.font, font_size=12, bold=True)  # Header font

# Add data rows
for row in data:
    row_cells = table.add_row().cells
    for i, column_name in enumerate(column_headers):
        row_cells[i].text = row.get(column_name, "")
        set_font(row_cells[i], args.font, font_size=11)  # Data font

# Save DOCX file
doc.save(args.targetpath)
print(f"Saved DOCX output to {args.targetpath}")
